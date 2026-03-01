#!/usr/bin/env python3
"""
Generate Sony & Nintendo Japanese Earnings Update Report (DOCX)
Report date: March 1, 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors ──────────────────────────────────────────────────────
SONY_BLUE   = RGBColor(0x00, 0x30, 0x87)
NINT_RED    = RGBColor(0xE6, 0x00, 0x12)
DARK_GRAY   = RGBColor(0x2C, 0x3E, 0x50)
MED_GRAY    = RGBColor(0x7F, 0x8C, 0x8D)
LIGHT_GRAY  = RGBColor(0xEC, 0xF0, 0xF1)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN       = RGBColor(0x27, 0xAE, 0x60)
RED_WARN    = RGBColor(0xC0, 0x39, 0x2B)
BLACK       = RGBColor(0x00, 0x00, 0x00)

def set_font(run, name="Times New Roman", size=10, bold=False,
             color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Japanese font
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Hiragino Sans')
    rPr.insert(0, rFonts)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, side='all', color='DDDDDD', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    sides = ['top','left','bottom','right'] if side == 'all' else [side]
    for s in sides:
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=DARK_GRAY, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_para(doc, text, size=10, color=BLACK, bold=False, indent=0, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, size=10, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run("■  " + text)
    set_font(run, size=size, color=color)
    return p

def add_chart(doc, path, caption, width_cm=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run2 = cap.add_run(caption)
    set_font(run2, size=8, italic=True, color=MED_GRAY)

def add_table_row(table, cells_data, header=False, bg_hex=None, bold=False):
    row = table.add_row()
    for i, (cell_text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(cell_text))
        set_font(run, size=9, bold=bold or header,
                 color=WHITE if (header or bg_hex == "003087") else BLACK)
        if bg_hex:
            set_cell_bg(cell, bg_hex)
        elif header:
            set_cell_bg(cell, "003087")
        set_cell_border(cell)
        cell.width = Cm(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

# ────────────────────────────────────────────────────────────────
# Build Document
# ────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ================================================================
# COVER / PAGE 1
# ================================================================
# Title banner
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(2)
r = p_title.add_run("決算アップデートレポート")
set_font(r, size=22, bold=True, color=DARK_GRAY)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(4)
r2 = p_sub.add_run("ソニーグループ (6758) ／ 任天堂 (7974) — 2025年12月期 第3四半期決算")
set_font(r2, size=13, bold=False, color=MED_GRAY)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_after = Pt(8)
r3 = p_date.add_run("作成日：2026年3月1日 ｜ 対象期間：2025年10月～12月（Q3）")
set_font(r3, size=9, italic=True, color=MED_GRAY)

add_horizontal_rule(doc)

# ── Summary boxes ───────────────────────────────────────────────
p_h = doc.add_paragraph()
p_h.paragraph_format.space_after = Pt(3)
r_h = p_h.add_run("■  決算サマリー")
set_font(r_h, size=12, bold=True, color=DARK_GRAY)

# Summary table (2 companies side by side)
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"

# --- Sony cell ---
sc = tbl.rows[0].cells[0]
sc.width = Cm(8.5)
set_cell_bg(sc, "EBF5FB")
set_cell_border(sc, color="003087", sz=8)
sp = sc.paragraphs[0]
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sc = sp.add_run("ソニーグループ（6758）")
set_font(r_sc, size=11, bold=True, color=SONY_BLUE)

for line in [
    ("決算期", "2025年12月 Q3（FY2025）"),
    ("発表日", "2026年2月5日"),
    ("レーティング", "OUTPERFORM（維持）"),
    ("目標株価", "¥3,500（変更なし）"),
    ("売上高", "¥3,714十億（前年比+1%）　▶ BEAT"),
    ("営業利益", "¥515十億（前年比+22%）　▶ BEAT"),
    ("営業利益率", "13.9%（前年比+240bps）"),
    ("純利益", "¥377十億（前年比+11%）"),
    ("EPS（ADR）", "$0.41（コンセンサス$0.33を24%超過）"),
    ("通期ガイダンス", "売上高¥12.3兆 / 営業利益¥1.54兆（上方修正）"),
]:
    p_row = sc.add_paragraph()
    p_row.paragraph_format.space_before = Pt(1)
    p_row.paragraph_format.space_after  = Pt(1)
    r_l = p_row.add_run(f"  {line[0]}：")
    set_font(r_l, size=9, bold=True, color=DARK_GRAY)
    r_v = p_row.add_run(line[1])
    set_font(r_v, size=9, color=BLACK)

# --- Nintendo cell ---
nc = tbl.rows[0].cells[1]
nc.width = Cm(8.5)
set_cell_bg(nc, "FDEDEC")
set_cell_border(nc, color="E60012", sz=8)
np_ = nc.paragraphs[0]
np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_nc = np_.add_run("任天堂（7974）")
set_font(r_nc, size=11, bold=True, color=NINT_RED)

for line in [
    ("決算期", "FY2026 Q3（2025年12月末）"),
    ("発表日", "2026年2月3日"),
    ("レーティング", "OUTPERFORM（維持）"),
    ("目標株価", "¥12,000（変更なし）"),
    ("Q3売上高", "¥758十億（前年比+86%）　▶ BEAT"),
    ("Q1-Q3累計売上", "¥1,906十億（前年比+99%）"),
    ("Q1-Q3営業利益", "¥300十億（前年比+21%）"),
    ("Switch 2販売", "累計1,737万台（Q3：701万台）"),
    ("デジタル売上", "¥126.5十億（前年比+60%）"),
    ("通期ガイダンス", "売上高¥2.25兆 / SW2 1,900万台（維持）"),
]:
    p_row = nc.add_paragraph()
    p_row.paragraph_format.space_before = Pt(1)
    p_row.paragraph_format.space_after  = Pt(1)
    r_l = p_row.add_run(f"  {line[0]}：")
    set_font(r_l, size=9, bold=True, color=DARK_GRAY)
    r_v = p_row.add_run(line[1])
    set_font(r_v, size=9, color=BLACK)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Key takeaways ────────────────────────────────────────────────
add_heading(doc, "主要なポイント", level=1, color=DARK_GRAY, size=12)

for bullet in [
    "ソニー：記録的な第3四半期営業利益を達成。I&SS（イメージセンサー）と音楽セグメントが牽引し、"
    "営業利益は前年比+22%の¥515十億円に急増。通期ガイダンスを売上高¥12.3兆・営業利益¥1.54兆に上方修正。",
    "任天堂：Nintendo Switch 2が歴史的な快走。発売から約6ヶ月で累計1,737万台を達成、第3四半期単体の"
    "売上は前年比+86%に急拡大。「マリオカート ワールド」は累計1,403万本のメガヒットを記録。",
    "両社比較：ソニーの営業利益率は13.9%、任天堂は19.2%と、いずれも大幅改善。"
    "ゲーム・エンタメ・テクノロジーの各領域で日本の両雄が収益力の高さを示した。",
    "リスク要因：ソニーはPS5後期サイクル・メモリコスト上昇、任天堂はSwitch 2の欧米需要見通し下振れ"
    "リスクを注視。いずれもコア事業の競争優位は健在であり、レーティング・目標株価を維持。",
]:
    add_bullet(doc, bullet, size=10)

doc.add_page_break()

# ================================================================
# PAGE 2-3: SONY DETAILED ANALYSIS
# ================================================================
add_heading(doc, "ソニーグループ（6758）— 詳細決算分析", level=1, color=SONY_BLUE, size=14)
add_para(doc, "Q3 FY2025（2025年10月1日〜12月31日） ｜ 発表日：2026年2月5日", size=9, color=MED_GRAY)
add_horizontal_rule(doc)

add_heading(doc, "1. 売上高分析", level=2, color=SONY_BLUE, size=12)
add_para(doc,
    "ソニーグループのQ3 FY2025における売上高は¥3,713.7十億円（前年比+1%）となり、"
    "市場コンセンサス¥3,680十億円を+¥33.7十億円（+0.9%）上回った。"
    "金融サービス子会社（ソニーフィナンシャルグループ）が2025年10月1日付で連結から除外されたため、"
    "前年同期比での単純比較には注意が必要だが、継続事業ベースでは堅調な成長を維持した。",
    size=10)

add_para(doc,
    "セグメント別では、ゲーム＆ネットワークサービス（G&NS）が売上¥1,613.6十億円（前年比-4%）と"
    "PS5ハードウェア販売減少の影響を受けた一方、ネットワークサービスや自社タイトルの好調により"
    "営業利益は¥140.8十億円（前年比+19%）と大幅増益。音楽セグメントは売上+13%、営業利益+9%と"
    "ライブエンタメ・ストリーミング双方が貢献。I&SSはスマートフォン市場回復と大口顧客"
    "（Apple等）向け出荷増加・単価改善により力強い伸びを示した。",
    size=10)

# Sony Revenue chart
add_chart(doc, "charts/sony_revenue.png",
          "図1：ソニーグループ 四半期別売上高推移（出所：ソニーグループ決算短信・当社推計）")

add_heading(doc, "2. 収益性分析（マージン）", level=2, color=SONY_BLUE, size=12)
add_para(doc,
    "Q3 FY2025の営業利益は¥515.0十億円（前年比+22%）、営業利益率は13.9%と前年11.5%から"
    "+240bpsの大幅改善を達成し、第3四半期として記録的な水準となった。純利益は¥377.3十億円"
    "（前年比+11%）。I&SSセグメントの高付加価値センサーへのミックスシフトと、"
    "音楽セグメントにおける「ピーナッツ・ホールディングス」追加取得に伴う約¥45十億円の"
    "再測定益が利益を押し上げた。一方、ソニー・ピクチャーズ（SPE）は前年の「ヴェノム：ザ・ラスト・ダンス」"
    "との比較が厳しく、売上-12%・営業利益-11%と低調だった。",
    size=10)

add_chart(doc, "charts/sony_operating_income.png",
          "図2：ソニーグループ 営業利益・営業利益率推移（出所：ソニーグループ決算短信・当社推計）")

# Segment table
add_heading(doc, "3. セグメント別実績サマリー（Q3 FY2025）", level=2, color=SONY_BLUE, size=12)

seg_table = doc.add_table(rows=1, cols=5)
seg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
seg_table.style = "Table Grid"

widths = [3.5, 2.5, 2.5, 2.5, 2.5]
headers = ["セグメント", "売上高（¥B）", "前年比", "営業利益（¥B）", "前年比"]
header_row = seg_table.rows[0]
for i, (h, w) in enumerate(zip(headers, widths)):
    cell = header_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, color=WHITE)
    set_cell_bg(cell, "003087")
    set_cell_border(cell)
    cell.width = Cm(w)

seg_data = [
    ("G&NS（ゲーム）", "1,613.6", "▼4%", "140.8", "▲19%"),
    ("I&SS（センサー）", "~585", "▲高成長", "~95", "▲高成長"),
    ("ET&S（エレクトロ）", "658.1", "▼7%", "59.4", "▼23%"),
    ("Music（音楽）", "542.3", "▲13%", "106.4", "▲9%"),
    ("Pictures（映像）", "~340", "▼12%", "~27", "▼11%"),
    ("合計（継続事業）", "3,713.7", "▲1%", "515.0", "▲22%"),
]
colors_alt = ["FFFFFF","EBF5FB","FFFFFF","EBF5FB","FFFFFF","D6EAF8"]
for data, bg in zip(seg_data, colors_alt):
    row = seg_table.add_row()
    bold_flag = data[0].startswith("合計")
    for i, (val, w) in enumerate(zip(data, widths)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        set_font(run, size=9, bold=bold_flag, color=BLACK)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.width = Cm(w)

cap_p = doc.add_paragraph()
cap_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cap_r = cap_p.add_run("出所：ソニーグループ決算短信 2026年2月5日　*I&SS・Picturesは円換算推計値")
set_font(cap_r, size=7, italic=True, color=MED_GRAY)

add_chart(doc, "charts/sony_segments.png",
          "図3：ソニーグループ セグメント別業績（Q3 FY2025）（出所：ソニーグループ決算短信 2026年2月5日）")
add_chart(doc, "charts/sony_beat_miss.png",
          "図4：ソニーQ3 FY2025 コンセンサス比（出所：Bloomberg / ソニーグループ決算短信）")

doc.add_page_break()

# ================================================================
# PAGE 4-5: SONY GUIDANCE + INVESTMENT THESIS
# ================================================================
add_heading(doc, "4. 通期ガイダンス修正（FY2025）", level=2, color=SONY_BLUE, size=12)
add_para(doc,
    "ソニーはQ3の好業績を受け、FY2025通期の業績予想を以下のとおり上方修正した。"
    "米国関税政策変更による約¥50十億円の下振れリスクを見込んでもなお大幅な上方修正となっており、"
    "経営陣の自信を示す内容だ。",
    size=10)

# Guidance table
gd_table = doc.add_table(rows=1, cols=4)
gd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
gd_table.style = "Table Grid"
gd_widths = [4.5, 2.8, 2.8, 2.8]
gd_headers = ["項目", "旧予想（11月時点）", "新予想（2月修正）", "修正幅"]
gd_header_row = gd_table.rows[0]
for i, (h, w) in enumerate(zip(gd_headers, gd_widths)):
    cell = gd_header_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, color=WHITE)
    set_cell_bg(cell, "003087")
    set_cell_border(cell)
    cell.width = Cm(w)

gd_data = [
    ("売上高（十億円）", "11,940", "12,300", "▲+3%"),
    ("営業利益（十億円）", "1,426", "1,540", "▲+8%"),
    ("純利益（十億円）", "—", "修正なし", "—"),
]
gd_colors = ["FFFFFF","EBF5FB","FFFFFF"]
for data, bg in zip(gd_data, gd_colors):
    row = gd_table.add_row()
    for i, (val, w) in enumerate(zip(data, gd_widths)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        clr = GREEN if "▲" in val else (RED_WARN if "▼" in val else BLACK)
        set_font(run, size=9, bold=(i==3), color=clr)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.width = Cm(w)

cap_p2 = doc.add_paragraph()
cap_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cap_r2 = cap_p2.add_run("出所：ソニーグループ 2026年2月5日決算発表資料")
set_font(cap_r2, size=7, italic=True, color=MED_GRAY)

add_chart(doc, "charts/sony_guidance.png",
          "図10：ソニー FY2025 通期業績予想修正（出所：ソニーグループ 2026年2月5日）")

add_heading(doc, "5. 投資テーゼへの影響", level=2, color=SONY_BLUE, size=12)

thesis_bullets_sony = [
    ("多角化ポートフォリオの強靭性 ▶ 維持・強化",
     "G&NSが減収のなか、I&SS・音楽・G&NS営業利益が補完し全体利益を押し上げた。"
     "異なる事業サイクルを持つポートフォリオのリスク分散効果が改めて実証された。"),
    ("I&SS（イメージセンサー）の競争優位 ▶ 強化",
     "スマートフォン市場回復とApple向け供給増、高解像度センサーへのミックス改善が"
     "売上単価・利益率を押し上げた。AI搭載デバイス普及に伴うセンサー需要増は中期的な追い風。"),
    ("PS5後期サイクルリスク ▶ 継続監視",
     "PS5ハードウェア販売は前年比減少。ただし月間アクティブユーザーは12月に1億3,200万人の"
     "過去最多を更新。ソフトウェア・PSNサービス収益への移行が着実に進んでいる点は評価できる。"),
    ("ソニー・フィナンシャルグループ分離の影響 ▶ 中立",
     "2025年10月の一部スピンオフにより連結構造が変化。継続事業ベースでの利益率は改善しており、"
     "事業集中戦略として肯定的に評価する。"),
]

for title, body in thesis_bullets_sony:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    r_t = p.add_run("■  " + title)
    set_font(r_t, size=10, bold=True, color=SONY_BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent  = Cm(0.7)
    p2.paragraph_format.space_after  = Pt(3)
    r_b = p2.add_run(body)
    set_font(r_b, size=9.5, color=BLACK)

doc.add_page_break()

# ================================================================
# PAGE 6-8: NINTENDO DETAILED ANALYSIS
# ================================================================
add_heading(doc, "任天堂（7974）— 詳細決算分析", level=1, color=NINT_RED, size=14)
add_para(doc, "Q3 FY2026（2025年10月1日〜12月31日） ｜ 発表日：2026年2月3日", size=9, color=MED_GRAY)
add_horizontal_rule(doc)

add_heading(doc, "6. 売上高分析（Q3単体 & 累計）", level=2, color=NINT_RED, size=12)
add_para(doc,
    "任天堂のQ3 FY2026（2025年10〜12月）単体売上高は前年比+86%と急拡大し、"
    "2025年6月に発売したNintendo Switch 2が本格的なホリデーシーズンを迎え"
    "過去最高水準に迫る¥758十億円規模の売上を達成した。"
    "Q1〜Q3累計では¥1,905.8十億円（前年比+99.3%）と、事実上の売上倍増を実現している。",
    size=10)

add_para(doc,
    "Nintendo Switch 2の成功を支えたのは、プレミアム価格設定（従来Switchより高い単価）と"
    "強力なソフトラインアップである。「マリオカート ワールド」が累計1,403万本、"
    "「ポケットモンスター レジェンズ：Z-A Nintendo Switch 2 Edition」が389万本と"
    "スマッシュヒットを記録した。一方、旧Switch（OLED/Lite）は前年比-66%の325万台と"
    "旧機種の終息が明確に進んでいる。",
    size=10)

add_chart(doc, "charts/nintendo_revenue.png",
          "図5：任天堂 四半期別売上高推移（出所：任天堂決算短信・当社推計）")

add_heading(doc, "7. Nintendo Switch 2 — 販売実績", level=2, color=NINT_RED, size=12)
add_para(doc,
    "Switch 2は2025年6月の発売からわずか6ヶ月強で累計1,737万台を達成し、"
    "任天堂史上最速ペースのハードウェア普及を記録している（Wii Uの生涯販売台数を既に超過）。"
    "Q3単体では701万台を販売し、ホリデー需要の強さを示した。"
    "任天堂は通期ガイダンスとして1,900万台を維持しており、Q4での280万台追加販売を見込む。",
    size=10)

# Switch 2 table
sw_table = doc.add_table(rows=1, cols=4)
sw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sw_table.style = "Table Grid"
sw_widths = [4.0, 2.8, 2.8, 2.8]
sw_headers = ["項目", "Q1 FY26", "Q2 FY26", "Q3 FY26（最新）"]
sw_header_row = sw_table.rows[0]
for i, (h, w) in enumerate(zip(sw_headers, sw_widths)):
    cell = sw_header_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, color=WHITE)
    set_cell_bg(cell, "E60012")
    set_cell_border(cell)
    cell.width = Cm(w)

sw_data = [
    ("Switch 2 四半期販売台数", "626万台", "410万台", "701万台"),
    ("Switch 2 累計販売台数", "626万台", "1,036万台", "1,737万台"),
    ("旧Switch 四半期販売台数", "約56万台", "約136万台", "約133万台"),
    ("ソフト販売（SW2）", "—", "1,590万本", "3,793万本（累計）"),
    ("デジタル売上（十億円）", "—", "—", "126.5十億（+60%YoY）"),
]
sw_colors_alt = ["FFFFFF","FDEDEC","FFFFFF","FDEDEC","FFFFFF"]
for data, bg in zip(sw_data, sw_colors_alt):
    row = sw_table.add_row()
    for i, (val, w) in enumerate(zip(data, sw_widths)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        set_font(run, size=9, color=BLACK)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.width = Cm(w)

cap_sw = doc.add_paragraph()
cap_sw.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cap_sw_r = cap_sw.add_run("出所：任天堂決算短信 2026年2月3日　*Q1/Q2はGoNintendo推計値を含む")
set_font(cap_sw_r, size=7, italic=True, color=MED_GRAY)

add_chart(doc, "charts/nintendo_switch2.png",
          "図6：Nintendo Switch 2 四半期別・累計販売台数（出所：任天堂決算短信 2026年2月3日）")

add_heading(doc, "8. 収益性分析", level=2, color=NINT_RED, size=12)
add_para(doc,
    "Q3単体の営業利益率は19.2%と回復したが、Q1〜Q3累計では15.8%（営業利益¥300.4十億円）にとどまった。"
    "Switch 2は従来機より高い製造コストを伴うため、売上増加に対して利益率の改善は緩やか。"
    "また、販売地域・製品ミックスが当初想定と異なり、日本国内販売比率の上昇や"
    "サードパーティソフトの比率増加がマージンを若干圧迫する要因となった。"
    "中長期的にはデジタル販売（ソフト売上の50.4%）比率の上昇が利益率を押し上げる見通し。",
    size=10)

add_chart(doc, "charts/nintendo_operating_income.png",
          "図7：任天堂 営業利益・営業利益率推移（出所：任天堂決算短信・当社推計）")

add_chart(doc, "charts/nintendo_software.png",
          "図8：Nintendo Switch 2 主要ソフトウェア販売本数（出所：任天堂決算短信 2026年2月3日）")

doc.add_page_break()

# ================================================================
# PAGE 9-10: NINTENDO THESIS + COMPARISON
# ================================================================
add_heading(doc, "9. 任天堂 投資テーゼへの影響", level=2, color=NINT_RED, size=12)

thesis_bullets_nint = [
    ("Nintendo Switch 2の牽引力 ▶ 大幅強化",
     "Switch 2は任天堂史上最速の普及ペースで1,737万台を達成。「マリオカート ワールド」の"
     "大ヒットと強力なホリデーラインアップが需要を支えた。プラットフォームとしての地位確立は"
     "中期的なソフト・サービス収益拡大の基盤となる。"),
    ("ソフトウェア・デジタルエコシステム ▶ 強化",
     "デジタル販売が前年比+60%と急成長し、ソフト売上の50.4%に到達。高マージンの"
     "デジタルビジネスへの移行が加速しており、プラットフォーム成熟とともに収益性改善が期待される。"),
    ("地域・製品ミックスリスク ▶ 継続監視",
     "欧米向け販売比率が想定より低く、日本・サードパーティソフト比率が高い方向にシフト。"
     "これが利益率にやや下押し圧力をかけた。欧米でのSwitch 2普及加速が今後の焦点となる。"),
    ("通期ガイダンスの達成確度 ▶ 高い",
     "売上高¥2.25兆・Switch 2販売1,900万台のガイダンスは維持されており、"
     "Q4での2月「マリオテニス フィーバー」・3月「ポケポキア」といった新タイトルも予定されている。"
     "需要の持続性は高く、ガイダンス達成の蓋然性は高いと判断する。"),
]

for title, body in thesis_bullets_nint:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    r_t = p.add_run("■  " + title)
    set_font(r_t, size=10, bold=True, color=NINT_RED)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent  = Cm(0.7)
    p2.paragraph_format.space_after  = Pt(3)
    r_b = p2.add_run(body)
    set_font(r_b, size=9.5, color=BLACK)

add_heading(doc, "10. ソニー vs 任天堂 比較分析", level=2, color=DARK_GRAY, size=12)

add_chart(doc, "charts/comparison_margins.png",
          "図9：ソニー vs 任天堂 営業利益率比較（出所：各社決算短信・当社推計）")

add_para(doc,
    "両社はいずれもQ3において高い収益性を示したが、事業ドライバーは対照的だ。"
    "ソニーはI&SS・音楽・G&NSという3本柱が補完し合う多角化モデルで安定的な利益成長を確保。"
    "任天堂はSwitch 2という単一プラットフォームへの集中によって爆発的な売上成長を実現したが、"
    "その分コスト構造の変化やミックスリスクへの感応度が高い。",
    size=10)

# Comparison table
comp_table = doc.add_table(rows=1, cols=3)
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
comp_table.style = "Table Grid"
comp_widths = [5.0, 3.5, 3.5]
comp_headers = ["指標（Q3 FY2025/26）", "ソニーグループ（6758）", "任天堂（7974）"]
comp_header_row = comp_table.rows[0]
for i, (h, w) in enumerate(zip(comp_headers, comp_widths)):
    cell = comp_header_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, color=WHITE)
    set_cell_bg(cell, "2C3E50")
    set_cell_border(cell)
    cell.width = Cm(w)

comp_data = [
    ("売上高（十億円）",           "3,714",           "758"),
    ("売上高 前年比",              "+1%",             "+86%"),
    ("営業利益（十億円）",          "515",             "143"),
    ("営業利益率",                 "13.9%",           "19.2%"),
    ("純利益（十億円）",            "377",             "約129"),
    ("コア成長ドライバー",          "センサー・音楽",   "Switch 2"),
    ("通期ガイダンス修正",          "上方修正（+8%）",  "変更なし（維持）"),
    ("レーティング",               "OUTPERFORM（維持）","OUTPERFORM（維持）"),
    ("目標株価",                   "¥3,500",          "¥12,000"),
]
comp_colors_alt = ["FFFFFF","F0F3F4","FFFFFF","F0F3F4","FFFFFF","F0F3F4","FFFFFF","F0F3F4","FFFFFF"]
for data, bg in zip(comp_data, comp_colors_alt):
    row = comp_table.add_row()
    for i, (val, w) in enumerate(zip(data, comp_widths)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        if "OUTPERFORM" in val:
            clr = GREEN
        elif "▲" in val or "上方" in val:
            clr = GREEN
        elif "▼" in val:
            clr = RED_WARN
        else:
            clr = BLACK
        set_font(run, size=9, bold=(i==0), color=clr)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.width = Cm(w)

cap_comp = doc.add_paragraph()
cap_comp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cap_comp_r = cap_comp.add_run("出所：各社決算短信・当社推計")
set_font(cap_comp_r, size=7, italic=True, color=MED_GRAY)

doc.add_page_break()

# ================================================================
# PAGE 11-12: VALUATION & ESTIMATES
# ================================================================
add_heading(doc, "バリュエーション・業績予想", level=1, color=DARK_GRAY, size=14)
add_horizontal_rule(doc)

add_heading(doc, "11. ソニーグループ 業績予想・バリュエーション", level=2, color=SONY_BLUE, size=12)

add_para(doc,
    "Q3結果とガイダンス上方修正を受け、ソニーのFY2025・FY2026予想を更新した。"
    "イメージセンサーの好調持続と音楽セグメントの安定成長を主要なアップサイドドライバーとして評価する。"
    "PS5サイクルの成熟はリスクとして注視しつつ、PSN/ゲームサービス収益の成長が代替する構造を評価。"
    "現在株価¥2,850（ADR：$21.40相当）に対して目標株価¥3,500は+23%のアップサイドを示す。"
    "OUTPERFORM（アウトパフォーム）を維持する。",
    size=10)

# Sony Estimates table
se_table = doc.add_table(rows=1, cols=5)
se_table.alignment = WD_TABLE_ALIGNMENT.CENTER
se_table.style = "Table Grid"
se_widths = [4.5, 2.2, 2.2, 2.2, 2.2]
se_headers = ["項目", "FY24実績", "FY25予想(旧)", "FY25予想(新)", "FY26予想(新)"]
se_header_row = se_table.rows[0]
for i, (h, w) in enumerate(zip(se_headers, se_widths)):
    cell = se_header_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, color=WHITE)
    set_cell_bg(cell, "003087")
    set_cell_border(cell)
    cell.width = Cm(w)

se_data = [
    ("売上高（兆円）", "11.89", "11.94", "12.30", "12.80"),
    ("売上高成長率", "+3.2%", "+0.4%", "+3.5%", "+4.1%"),
    ("営業利益（千億円）", "12.0", "14.26", "15.40", "17.00"),
    ("営業利益率", "10.1%", "11.9%", "12.5%", "13.3%"),
    ("純利益（千億円）", "9.1", "—", "~11.5", "~13.0"),
    ("EPS（ADR $）", "1.52", "—", "1.85", "2.15"),
    ("P/E（倍）", "14.1x", "—", "11.6x", "9.9x"),
    ("EV/EBITDA（倍）", "8.2x", "—", "7.1x", "6.5x"),
]
se_colors_alt = ["FFFFFF","EBF5FB"]*4
for data, bg in zip(se_data, se_colors_alt):
    row = se_table.add_row()
    for i, (val, w) in enumerate(zip(data, se_widths)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        set_font(run, size=9, bold=(i==3), color=SONY_BLUE if i==3 else BLACK)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.width = Cm(w)

cap_se = doc.add_paragraph()
cap_se.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cap_se_r = cap_se.add_run("出所：当社推計、ソニーグループ決算発表資料。E=Estimate（推計値）、A=Actual（実績）")
set_font(cap_se_r, size=7, italic=True, color=MED_GRAY)

add_heading(doc, "12. 任天堂 業績予想・バリュエーション", level=2, color=NINT_RED, size=12)

add_para(doc,
    "Switch 2の好調な立ち上がりを踏まえ、任天堂のFY2026通期予想は会社ガイダンス"
    "（売上高¥2.25兆、Switch 2販売1,900万台）に沿った水準を想定する。"
    "FY2027はデジタル販売比率の一段の上昇とサブスクリプションサービス（NSO）拡大により"
    "利益率改善が加速する見込み。目標株価¥12,000は予想EPS¥500のPER24倍に相当し、"
    "高いプラットフォームブランド価値と強固なIP資産に対するプレミアムを反映している。",
    size=10)

# Nintendo Estimates table
ne_table = doc.add_table(rows=1, cols=5)
ne_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ne_table.style = "Table Grid"
ne_widths = [4.5, 2.2, 2.2, 2.2, 2.2]
ne_headers = ["項目", "FY25実績", "FY26会社予想", "FY26当社予想", "FY27当社予想"]
ne_header_row = ne_table.rows[0]
for i, (h, w) in enumerate(zip(ne_headers, ne_widths)):
    cell = ne_header_row.cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, color=WHITE)
    set_cell_bg(cell, "E60012")
    set_cell_border(cell)
    cell.width = Cm(w)

ne_data = [
    ("売上高（兆円）", "1.19", "2.25", "2.25", "2.10"),
    ("売上高成長率", "+1%", "+89%", "+89%", "▼7%"),
    ("営業利益（千億円）", "5.8", "—", "~5.2", "~5.8"),
    ("営業利益率", "48.6%*", "—", "23.1%", "27.6%"),
    ("純利益（千億円）", "5.2", "—", "~5.5", "~5.9"),
    ("EPS（円）", "435", "—", "460", "500"),
    ("P/E（倍）", "22.2x", "—", "20.9x", "19.2x"),
    ("Switch 2販売（百万台）", "—", "19.0M", "19.0M", "20.0M"),
]
ne_colors_alt = ["FFFFFF","FDEDEC"]*4
for data, bg in zip(ne_data, ne_colors_alt):
    row = ne_table.add_row()
    for i, (val, w) in enumerate(zip(data, ne_widths)):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        set_font(run, size=9, bold=(i==3), color=NINT_RED if i==3 else BLACK)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.width = Cm(w)

cap_ne = doc.add_paragraph()
cap_ne.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cap_ne_r = cap_ne.add_run("出所：当社推計、任天堂決算発表資料。*FY25営業利益率は旧Switch成熟期の高水準を反映")
set_font(cap_ne_r, size=7, italic=True, color=MED_GRAY)

doc.add_page_break()

# ================================================================
# PAGE 13: SOURCES & DISCLAIMER
# ================================================================
add_heading(doc, "出典・免責事項", level=1, color=DARK_GRAY, size=12)
add_horizontal_rule(doc)

add_heading(doc, "参照資料（ソニーグループ）", level=2, color=SONY_BLUE, size=11)
sources_sony = [
    "ソニーグループ 2025年度第3四半期決算短信（発表日：2026年2月5日）",
    "  → https://www.sony.com/ja/SonyInfo/IR/library/fr.html",
    "ソニーグループ Q3 FY2025 決算説明会スライド（2026年2月5日）",
    "  → https://investgame.net/wp-content/uploads/2026/02/25q3_sonypre.pdf",
    "Sony Addict - Q3 FY2025 Consolidated Financial Results（2026年2月6日）",
    "  → https://sonyaddict.com/2026/02/06/sony-q3-fy2025-consolidated-financial-results/",
    "Variety - Sony Pictures Q3 2025 Results（2026年）",
    "  → https://variety.com/2026/film/asia/sony-pictures-revenue-sony-group-earnings-q3-2025-1236652653/",
    "Investing.com - Sony Q3 FY2025 Operating Income Jump（2026年）",
    "  → https://www.investing.com/news/company-news/sony-q3-fy2025-slides-22-operating-income-jump-prompts-forecast-upgrade-93CH-4486713",
]
for s in sources_sony:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if s.startswith("  →"):
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(s)
        set_font(run, size=8, italic=True, color=SONY_BLUE)
    else:
        run = p.add_run("• " + s)
        set_font(run, size=9, color=DARK_GRAY)

add_heading(doc, "参照資料（任天堂）", level=2, color=NINT_RED, size=11)
sources_nint = [
    "任天堂 2026年3月期 第3四半期決算短信（発表日：2026年2月3日）",
    "  → https://www.nintendo.co.jp/ir/en/library/earnings/index.html",
    "任天堂 Q3 FY2026 Financial Results Explanatory Material（2026年2月3日）",
    "  → https://www.nintendo.co.jp/ir/pdf/2026/260203_3e.pdf",
    "CNBC - Nintendo Q3 Earnings: Switch 2 Forecast（2026年2月3日）",
    "  → https://www.cnbc.com/2026/02/03/nintendo-q3-earnings-switch-2-forecast.html",
    "Nintendo Everything - Nintendo Financial Results February 2026",
    "  → https://nintendoeverything.com/nintendo-financial-results-february-2026-switch-2-at-17-37-million-units-switch-at-155-37-million-more/",
    "GoNintendo - Q3 FY2026 Results（2026年）",
    "  → https://gonintendo.com/contents/57281-nintendo-results-for-q3-of-fiscal-year-ending-march-2026-switch-at-155-37-million-switch-2-at-17-37-million",
]
for s in sources_nint:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if s.startswith("  →"):
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(s)
        set_font(run, size=8, italic=True, color=NINT_RED)
    else:
        run = p.add_run("• " + s)
        set_font(run, size=9, color=DARK_GRAY)

add_heading(doc, "コンセンサス推計源", level=2, color=DARK_GRAY, size=11)
add_para(doc,
    "• Bloomberg コンセンサス推計（決算発表前日2026年2月4日時点）\n"
    "• FactSet コンセンサス推計（2026年2月3日時点）\n"
    "• 当社アナリスト推計（独自モデル、2026年2月末更新）",
    size=9, color=DARK_GRAY)

add_horizontal_rule(doc)

add_heading(doc, "免責事項", level=2, color=MED_GRAY, size=10)
add_para(doc,
    "本レポートは情報提供のみを目的として作成されており、特定の有価証券の売買を推奨・勧誘するものではありません。"
    "本レポートに記載された情報は信頼できると考えられる情報源に基づいておりますが、その正確性・完全性を保証するもの"
    "ではありません。将来の業績・株価等に関する記述は見通しに基づくものであり、実際の結果とは異なる場合があります。"
    "投資判断は最終的にご自身の責任において行ってください。本資料の無断転載・引用を禁じます。\n"
    "作成日：2026年3月1日",
    size=8, color=MED_GRAY)

# Save
output_path = "Sony_Nintendo_Q3FY2025_Earnings_Update_JP.docx"
doc.save(output_path)
print(f"\n✅ Report saved: {output_path}")
